"""
Resume text extraction. Pulls raw text out of PDF and DOCX files.
This is the first stage of the parsing pipeline — structuring happens
in resume_structurer.py.
"""
import io

import pdfplumber
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text out of tables, since some resumes use table layouts.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def extract_text(file_bytes: bytes, file_type: str) -> str:
    file_type = file_type.lower().lstrip(".")
    if file_type == "pdf":
        return extract_text_from_pdf(file_bytes)
    if file_type == "docx":
        return extract_text_from_docx(file_bytes)
    raise ValueError(f"Unsupported file type: {file_type}")
